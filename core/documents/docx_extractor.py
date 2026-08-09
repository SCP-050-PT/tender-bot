"""
core/documents/docx_extractor.py
Извлечение текста из DOCX-файлов.
Вынесено из document_processor.py (v6.5).
ИСПРАВЛЕНО (v6.7.3):
  - Исправлен вызов self._extract_pdf_from_zip в _is_valid_docx (функция → метод класса)
  - _is_valid_docx теперь метод класса DocxExtractor
"""

import re
import zipfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from loguru import logger


try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx не установлен, DOCX через zipfile")

MAX_DOCX_TIMEOUT = 60
# Magic bytes
DOC_MAGIC = b"\xd0\xcf\x11\xe0"
DOCX_MAGIC = b"PK\x03\x04"
DOCX_MAGIC_ALT = b"PK\x05\x06"

QUANTITY_KEYWORDS = [
    "количество", "кол-во", "кол.", "объем", "объём",
    "планируемое кол-во", "кол-во обучаемых", "обучаемых",
    "чел.", "человек", "слушателей",
    "количество рабочих мест", "кол-во рабочих мест",
]


def _detect_real_format(file_path: Path) -> str:
    """Определяет реальный формат файла по magic bytes."""
    try:
        with open(file_path, "rb") as f:
            header = f.read(8)
        if header[:4] == DOC_MAGIC:
            return "doc"
        if header[:4] == DOCX_MAGIC or header[:4] == DOCX_MAGIC_ALT:
            return "docx"
        if header[:2] == b"PK":
            return "zip_like"
        return "unknown"
    except Exception as e:
        logger.warning(f"Не удалось определить формат {file_path}: {e}")
        return "unknown"


class DocxExtractor:
    """Извлекает текст из DOCX-файлов с заголовками таблиц."""

    def extract(self, file_path: Path, doc_name: str, force_docx: bool = False) -> str:
        """Извлекает текст из DOCX с таймаутом и проверкой формата."""
        real_format = _detect_real_format(file_path)
        logger.debug(f"Файл {doc_name}: реальный формат={real_format}, путь={file_path}")

        # Если .doc но на самом деле DOCX — переименуем
        if str(file_path).lower().endswith(".doc") and real_format == "docx":
            new_path = Path(str(file_path) + ".docx")
            try:
                import shutil
                shutil.copy2(file_path, new_path)
                file_path = new_path
                logger.info(f"[DocxExtractor] Переименован .doc → .docx: {doc_name}")
            except Exception as e:
                logger.warning(f"Не удалось переименовать: {e}")

        # Старый .doc — другой экстрактор
        if real_format == "doc":
            logger.warning(f"[DocxExtractor] Старый формат .doc: {doc_name}. Пробуем fallback.")
            return self._extract_old_doc(file_path)

        # Проверка валидности DOCX
        if not self._is_valid_docx(file_path, doc_name):
            logger.warning(f"[DocxExtractor] Невалидный DOCX: {doc_name}")
            pdf_text = self._extract_pdf_from_zip(file_path)
            if pdf_text:
                logger.info(f"[DocxExtractor] Извлечён текст из PDF в ZIP: {doc_name}")
                return pdf_text
            return ""

        def _do_extract():
            try:
                if HAS_DOCX:
                    return self._extract_with_docx(file_path)
                return self._extract_via_zip(file_path)
            except Exception as e:
                logger.error(f"Ошибка извлечения DOCX: {e}")
                return ""

        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(_do_extract)
                return future.result(timeout=MAX_DOCX_TIMEOUT)
        except FutureTimeoutError:
            logger.error(f"Таймаут {MAX_DOCX_TIMEOUT}с на DOCX: {doc_name}")
            zip_text = self._extract_via_zip(file_path)
            if zip_text:
                logger.info(f"[DocxExtractor] Fallback zipfile после таймаута: {len(zip_text)} симв.")
                return zip_text
            return "[ТАЙМАУТ: файл слишком сложный для обработки]"
        except Exception as e:
            logger.error(f"Ошибка при таймауте: {e}")
            return ""

    def _is_valid_docx(self, file_path: Path, doc_name: str = "") -> bool:
        """Проверяет, является ли файл валидным DOCX (ZIP с word/document.xml)."""
        fmt = _detect_real_format(file_path)
        if fmt not in ("docx", "zip_like"):
            return False

        try:
            with zipfile.ZipFile(file_path, "r") as z:
                namelist = z.namelist()
                has_word_doc = "word/document.xml" in namelist
                has_docprops = "docProps/core.xml" in namelist

                if not has_word_doc:
                    if "[Content_Types].xml" in namelist:
                        with z.open("[Content_Types].xml") as ct:
                            ct_content = ct.read().decode("utf-8", errors="ignore")
                            if "themeManager" in ct_content and "wordprocessingml" not in ct_content:
                                logger.warning(f"[DocxExtractor] Файл является theme/template, не документом")
                                # Пробуем извлечь PDF из themeManager-архива
                                pdf_text = self._extract_pdf_from_zip(file_path)
                                if pdf_text:
                                    logger.info(f"[DocxExtractor] Извлечён PDF из themeManager: {doc_name}")
                                    return True
                    return has_word_doc or has_docprops
        except zipfile.BadZipFile:
            return False
        except Exception as e:
            logger.debug(f"Ошибка проверки DOCX: {e}")
            return False

        return True

    def _extract_pdf_from_zip(self, file_path: Path) -> str:
        """Ищет PDF внутри ZIP-архива и извлекает текст."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                namelist = z.namelist()
                pdf_files = [n for n in namelist if n.lower().endswith(".pdf")]
                if pdf_files:
                    pdf_name = pdf_files[0]
                    with z.open(pdf_name) as f:
                        pdf_data = f.read()
                        import tempfile
                        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                            tmp.write(pdf_data)
                            tmp_path = tmp.name
                        try:
                            return self._extract_pdf_text(tmp_path)
                        finally:
                            Path(tmp_path).unlink(missing_ok=True)
        except Exception as e:
            logger.debug(f"PDF из ZIP ошибка: {e}")
        return ""

    def _extract_pdf_text(self, pdf_path: str) -> str:
        """Извлекает текст из PDF через pdfplumber или PyPDF2."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts)
        except ImportError:
            logger.debug("pdfplumber не установлен")
        except Exception as e:
            logger.debug(f"pdfplumber ошибка: {e}")

        try:
            import PyPDF2
            text_parts = []
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts)
        except ImportError:
            logger.debug("PyPDF2 не установлен")
        except Exception as e:
            logger.debug(f"PyPDF2 ошибка: {e}")

        return ""

    def _extract_old_doc(self, file_path: Path) -> str:
        """Fallback для старых .doc файлов (MS Compound Document)."""
        # antiword
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", str(file_path)], capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout:
                return result.stdout
        except FileNotFoundError:
            logger.debug("antiword не установлен")
        except Exception as e:
            logger.debug(f"antiword ошибка: {e}")

        # textract
        try:
            import textract
            text = textract.process(str(file_path)).decode("utf-8", errors="ignore")
            return text
        except ImportError:
            logger.debug("textract не установлен")
        except Exception as e:
            logger.debug(f"textract ошибка: {e}")

        # Бинарное извлечение UTF-16LE
        try:
            with open(file_path, "rb") as f:
                data = f.read()
            text_parts = []
            i = 0
            while i < len(data) - 1:
                if data[i] >= 0x20 and data[i] <= 0x7E and data[i + 1] == 0x00:
                    j = i
                    while j < len(data) - 1 and data[j + 1] == 0x00 and data[j] >= 0x20:
                        j += 2
                    chunk = data[i:j].decode("utf-16le", errors="ignore")
                    if len(chunk) > 5:
                        text_parts.append(chunk)
                    i = j
                else:
                    i += 1
            if text_parts:
                return " ".join(text_parts)
        except Exception as e:
            logger.debug(f"Бинарное извлечение ошибка: {e}")

        logger.error(f"Не удалось извлечь текст из .doc: {file_path}")
        return ""

    def _extract_with_docx(self, file_path: Path) -> str:
        """Извлекает через python-docx с заголовками таблиц."""
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        tables_text = []
        for table in doc.tables:
            try:
                headers = self._extract_headers(table)
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0 and headers:
                        continue
                    try:
                        row_text = self._format_row_with_headers(row, headers)
                        if row_text:
                            tables_text.append(" | ".join(row_text))
                    except Exception as e:
                        logger.debug(f"Ошибка чтения строки: {e}")
                        continue
            except Exception as e:
                logger.debug(f"Ошибка чтения таблицы: {e}")
                continue

        return "\n".join(paragraphs + tables_text)

    def _extract_headers(self, table) -> list:
        """Извлекает заголовки из первой строки таблицы."""
        if not table.rows:
            return []
        try:
            return [cell.text.strip().lower() for cell in table.rows[0].cells]
        except Exception as e:
            logger.debug(f"Ошибка чтения заголовков: {e}")
            return []

    def _format_row_with_headers(self, row, headers: list) -> list:
        """Форматирует строку таблицы с подписями заголовков."""
        result = []
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if not cell_text:
                continue
            if col_idx < len(headers) and headers[col_idx]:
                header = headers[col_idx]
                if self._looks_like_quantity(header, cell_text):
                    result.append(f"{header}: {cell_text}")
                else:
                    result.append(cell_text)
            else:
                result.append(cell_text)
        return result

    def _looks_like_quantity(self, header: str, cell_text: str) -> bool:
        """Проверяет, похоже ли содержимое на количество по заголовку."""
        if not cell_text:
            return False
        if not re.match(r"^[\d\s.,]+$", cell_text.replace(" ", "")):
            return False
        return any(kw in header.lower() for kw in QUANTITY_KEYWORDS)

    def _extract_via_zip(self, file_path: Path) -> str:
        """Fallback: извлечение через zipfile с проверкой структуры."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                namelist = z.namelist()
                doc_xml_path = None
                for name in namelist:
                    if name.endswith("word/document.xml") or name == "word/document.xml":
                        doc_xml_path = name
                        break

                if not doc_xml_path:
                    xml_files = [n for n in namelist if n.endswith(".xml")]
                    service_xml = [
                        "[Content_Types].xml", "_rels/.rels",
                        "docProps/core.xml", "docProps/app.xml",
                    ]
                    content_xml = [n for n in xml_files if n not in service_xml]
                    if content_xml:
                        word_xml = [n for n in content_xml if "word/" in n]
                        if word_xml:
                            doc_xml_path = word_xml[0]
                        else:
                            doc_xml_path = content_xml[0]
                        logger.warning(f"word/document.xml не найден, используем {doc_xml_path}")
                    else:
                        logger.error(f"Нет XML-файлов в архиве: {namelist[:10]}")
                        return ""

                with z.open(doc_xml_path) as f:
                    xml_content = f.read().decode("utf-8", errors="ignore")
                    text = re.sub(r"<[^>]+>", "", xml_content)
                    text = re.sub(r"<", "<", text)
                    text = re.sub(r">", ">", text)
                    text = re.sub(r"&", "&", text)
                    text = re.sub(r""", '"', text)
                    text = re.sub(r"'", "'", text)
                    text = re.sub(r"\s+", " ", text).strip()
                    return text
        except zipfile.BadZipFile:
            logger.error(f"Файл не является ZIP-архивом: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Ошибка zip-извлечения: {e}")
            return ""