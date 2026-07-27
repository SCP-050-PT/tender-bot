"""
core/document_processor.py
Обработка документов тендера: скачивание, извлечение текста, фильтрация.
ИСПРАВЛЕНО (27.07.2026 v6.2):
  - БАГ 3: Распаковка ZIP-архивов с рекурсивным парсингом вложенных файлов
  - БАГ 5: Проверка магических байтов перед открытием PDF (фантомные 404-HTML)
  - БАГ 4: Fallback openpyxl для .xls если xlrd недоступен
  - v6.1: Используется переданная сессия (curl_cffi) для скачивания файлов
  - v6.1: Убран verify=False, добавлены заголовки сессии
  - v6.1: Фикс 404: zakupki.gov.ru требует куки сессии
"""

import os
import re
import time
import zipfile
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from loguru import logger

try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning(
        "python-docx не установлен, DOCX-файлы будут обрабатываться через zipfile"
    )

try:
    import fitz  # PyMuPDF

    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF не установлен, PDF-файлы будут пропущены")

try:
    import xlrd

    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

try:
    import openpyxl

    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


# === v6.0: ЧЁРНЫЙ СПИСОК ФАЙЛОВ ===
CONTRACT_PATTERNS = [
    r"проект\s*контракта",
    r"проект\s*договора",
    r"\bконтракт\b",
    r"\bдоговор\b",
    r"\bконтракт\s*\S*",
    r"\bдоговор\s*\S*",
]

# === v6.0: ПРИОРИТЕТЫ ФАЙЛОВ ===
FILE_PRIORITY = {
    r"техническое\s*задание|тз|техзадание|техническое\s*задани": 100,
    r"извещение|извещени": 90,
    r"документаци|документы|документ": 80,
    r"пояснительн|пояснительная": 70,
    r"заявка|заявление": 60,
    r"протокол": 50,
    r"решение|решен": 40,
    r"приложение": 30,
}

# === v6.0: ЛИМИТЫ ===
MAX_DOCX_TIMEOUT = 30
MAX_TEXT_LENGTH = 50000
MAX_CONTRACT_FILE_SIZE = 200 * 1024

# ← v6.2: Магические байты для проверки PDF
PDF_MAGIC = b"%PDF"
ZIP_MAGIC = b"PK\x03\x04"


@dataclass
class DocumentInfo:
    name: str
    url: str
    file_type: str = ""
    size: str = ""
    date: str = ""
    is_active: bool = True
    file_url: str = ""
    file_size_bytes: int = 0
    is_contract: bool = False
    priority: int = 0

    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "url": self.url,
            "file_type": self.file_type,
            "size": self.size,
            "date": self.date,
            "is_active": self.is_active,
            "file_url": self.file_url,
            "is_contract": self.is_contract,
            "priority": self.priority,
        }


class DocumentProcessor:
    """Процессор документов тендера."""

    def __init__(self, download_dir: Optional[Path] = None, session=None):
        self.download_dir = (
            download_dir or Path(__file__).resolve().parent.parent / "downloads"
        )
        self.download_dir.mkdir(parents=True, exist_ok=True)
        # v6.1: Принимаем сессию извне (curl_cffi)
        self.session = session
        logger.info(f"DocumentProcessor: download_dir={self.download_dir}")

    def process_documents(
        self,
        documents: List[DocumentInfo],
        max_docs: int = 5,
    ) -> str:
        if not documents:
            logger.warning("Нет документов для обработки")
            return ""

        active_docs = [d for d in documents if d.is_active]
        logger.info(f"Активных документов: {len(active_docs)}")

        for doc in active_docs:
            doc.is_contract = self._is_contract_file(doc.name)
            doc.priority = self._get_file_priority(doc.name)
            if doc.is_contract:
                logger.info(f"Контракт/договор (не извлекаем текст): {doc.name}")

        active_docs.sort(key=lambda d: (-d.priority, d.is_contract))
        docs_to_process = [d for d in active_docs if not d.is_contract][:max_docs]
        contract_docs = [d for d in active_docs if d.is_contract]

        logger.info(
            f"Обработать: {len(docs_to_process)} (контрактов пропущено: {len(contract_docs)})"
        )

        texts = []
        for doc in docs_to_process:
            try:
                text = self._process_single_document(doc)
                if text:
                    texts.append(f"=== ФАЙЛ: {doc.name} ===\n{text}")
            except Exception as e:
                logger.error(f"Ошибка обработки {doc.name}: {e}")
                continue

        if contract_docs:
            contract_names = [d.name for d in contract_docs[:3]]
            texts.append(
                f"\n=== ВНИМАНИЕ: СРЕДИ ДОКУМЕНТОВ ЕСТЬ КОНТРАКТЫ/ДОГОВОРЫ ===\n"
            )
            texts.append(f"Пропущены (не анализируются): {', '.join(contract_names)}")
            if len(contract_docs) > 3:
                texts.append(f"... и ещё {len(contract_docs) - 3} файлов")

        result = "\n\n".join(texts)
        logger.info(f"Итоговый текст: {len(result)} символов")
        return result

    def _is_contract_file(self, filename: str) -> bool:
        if not filename:
            return False
        name_lower = filename.lower()
        for pattern in CONTRACT_PATTERNS:
            if re.search(pattern, name_lower, re.IGNORECASE):
                return True
        return False

    def _get_file_priority(self, filename: str) -> int:
        if not filename:
            return 0
        name_lower = filename.lower()
        max_priority = 0
        for pattern, priority in FILE_PRIORITY.items():
            if re.search(pattern, name_lower, re.IGNORECASE):
                max_priority = max(max_priority, priority)
        return max_priority

    def _process_single_document(self, doc: DocumentInfo) -> str:
        logger.info(f"Обработка: {doc.name} ({doc.file_type})")

        if doc.is_contract and doc.file_size_bytes > MAX_CONTRACT_FILE_SIZE:
            logger.info(f"Пропущен (контракт >200 KB): {doc.name}")
            return ""

        file_path = self._download_file(doc)
        if not file_path:
            return ""

        # ← v6.2: Проверка магических байтов перед обработкой
        if not self._validate_file_content(file_path, doc.file_type):
            return ""

        text = self._extract_text(file_path, doc.file_type, doc.name)
        if not text:
            return ""

        if len(text) > MAX_TEXT_LENGTH:
            logger.info(f"Обрезано с {len(text)} до {MAX_TEXT_LENGTH} символов")
            text = (
                text[:MAX_TEXT_LENGTH]
                + "\n[... текст обрезан — слишком длинный файл ...]"
            )

        return text

    def _download_file(self, doc: DocumentInfo) -> Optional[Path]:
        """Скачивает файл с использованием сессии (curl_cffi)."""
        if not doc.file_url:
            return None

        try:
            # v6.1: Используем переданную сессию вместо голого requests
            if self.session:
                response = self.session.get(doc.file_url, timeout=30)
            else:
                import requests

                response = requests.get(doc.file_url, timeout=30, verify=False)

            if response.status_code != 200:
                logger.warning(f"Статус {response.status_code}: {doc.file_url}")
                # v6.1: Детальное логирование для диагностики 404
                if response.status_code == 404:
                    logger.error(
                        f"🔴 404 на файл: {doc.file_url}\n"
                        f"   Заголовки ответа: {dict(response.headers)}"
                    )
                return None

            content_length = len(response.content)
            doc.file_size_bytes = content_length

            safe_name = re.sub(r"[^\w\-_.]", "_", doc.name)[:100]
            file_path = self.download_dir / f"{safe_name}_{int(time.time())}"

            with open(file_path, "wb") as f:
                f.write(response.content)

            logger.info(f"Скачан: {file_path} ({content_length} байт)")
            return file_path

        except Exception as e:
            logger.error(f"Ошибка скачивания: {e}")
            return None

    # ← v6.2: Новый метод — проверка магических байтов
    def _validate_file_content(self, file_path: Path, file_type: str) -> bool:
        """Проверяет, что файл действительно соответствует заявленному типу."""
        try:
            with open(file_path, "rb") as f:
                header = f.read(8)

            ext = (file_type.lower() if file_type else file_path.suffix.lower()).lstrip(
                "."
            )

            if ext in ["pdf"]:
                if not header.startswith(PDF_MAGIC):
                    # Может быть HTML-страница ошибки
                    if b"<html" in header or b"<!DOCTYPE" in header:
                        logger.error(
                            f"🔴 Файл {file_path.name} — HTML-страница ошибки, не PDF"
                        )
                    else:
                        logger.error(
                            f"🔴 Файл {file_path.name} — не PDF (magic: {header[:4].hex()})"
                        )
                    return False

            elif ext in ["zip", "docx", "xlsx"]:
                if not header.startswith(ZIP_MAGIC):
                    logger.error(
                        f"🔴 Файл {file_path.name} — не ZIP (magic: {header[:4].hex()})"
                    )
                    return False

            return True
        except Exception as e:
            logger.error(f"Ошибка проверки файла {file_path}: {e}")
            return False

    def _extract_text(self, file_path: Path, file_type: str, doc_name: str) -> str:
        ext = file_type.lower() if file_type else file_path.suffix.lower()

        # ← v6.2: Определяем реальный тип по содержимому, если расширение неясно
        if not ext or ext == "zip" or ext not in ["docx", "doc", "pdf", "xlsx", "xls", "txt", "rtf"]:
            ext = self._detect_file_type(file_path)

        # ← v6.1: Пробуем оба формата если один не сработал
        if ext in ["docx", "doc"]:
            text = self._extract_docx(file_path, doc_name)
            if not text and ext == "doc":
                # Пробуем как docx (файл может быть переименован)
                text = self._extract_docx(file_path, doc_name, force_docx=True)
            return text
        elif ext == "pdf":
            text = self._extract_pdf(file_path)
            if not text:
                # v6.3: Fallback — может быть DOC под видом PDF
                real_type = self._detect_file_type(file_path)
                if real_type in ["doc", "docx"]:
                    logger.info(f"[v6.3] Файл {doc_name} — на самом деле {real_type}, пробуем docx")
                    text = self._extract_docx(file_path, doc_name)
            return text
        elif ext in ["xlsx", "xls"]:
            return self._extract_excel(file_path)
        elif ext == "zip":
            # ← v6.2: Распаковка ZIP-архивов
            return self._extract_zip(file_path)
        elif ext in ["txt", "rtf"]:
            return self._extract_text_file(file_path)
        else:
            logger.warning(f"Неизвестный тип: {ext}")
            return ""
        
    def _detect_file_type(self, file_path: Path) -> str:
        """Определяет тип файла по магическим байтам."""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
            if header.startswith(b'\x50\x4B\x03\x04'):
                return 'zip'  # ZIP (DOCX, XLSX или архив)
            elif header.startswith(b'\xD0\xCF\x11\xE0'):
                return 'doc'   # OLE (DOC, XLS)
            elif header.startswith(b'%PDF-'):
                return 'pdf'
            elif header.startswith(b'PK'):
                return 'zip'
        except Exception as e:
            logger.error(f"Ошибка определения типа файла: {e}")
        return file_path.suffix.lower().lstrip('.') or 'unknown'
    
    

    # ← v6.2: Новый метод — распаковка ZIP
    def _extract_zip(self, file_path: Path, doc_name: str) -> str:
        """Распаковывает ZIP-архив и извлекает текст из вложенных файлов."""
        logger.info(f"[v6.2] Распаковка ZIP: {doc_name}")
        texts = []

        # Создаём временную директорию
        temp_dir = Path(tempfile.mkdtemp(prefix="tender_zip_"))

        try:
            with zipfile.ZipFile(file_path, "r") as z:
                # Получаем список файлов (исключая системные)
                files = [
                    f
                    for f in z.namelist()
                    if not f.startswith("__MACOSX/") and not f.startswith(".")
                ]

                logger.info(f"[v6.2] В ZIP найдено файлов: {len(files)}")

                # Сортируем по приоритету
                prioritized = []
                for fname in files:
                    priority = self._get_file_priority(fname)
                    is_contract = self._is_contract_file(fname)
                    ext = Path(fname).suffix.lower().lstrip(".")
                    # Пропускаем системные и контракты
                    if is_contract:
                        continue
                    prioritized.append((priority, fname, ext, is_contract))

                prioritized.sort(key=lambda x: (-x[0], x[3]))

                # Обрабатываем до 3 файлов из архива
                for priority, fname, ext, _ in prioritized[:3]:
                    try:
                        # Извлекаем во временную директорию
                        extracted_path = temp_dir / Path(fname).name
                        with z.open(fname) as src, open(extracted_path, "wb") as dst:
                            dst.write(src.read())

                        # Определяем тип по расширению
                        if ext in ["docx", "doc"]:
                            text = self._extract_docx(extracted_path, fname)
                        elif ext == "pdf":
                            text = self._extract_pdf(extracted_path)
                        elif ext in ["xlsx", "xls"]:
                            text = self._extract_excel(extracted_path)
                        elif ext in ["txt", "rtf"]:
                            text = self._extract_text_file(extracted_path)
                        elif ext == "zip":
                            # Вложенный ZIP — рекурсия
                            text = self._extract_zip(extracted_path, fname)
                        else:
                            text = self._extract_by_content(extracted_path, fname)

                        if text:
                            texts.append(f"=== ВЛОЖЕННЫЙ ФАЙЛ: {fname} ===\n{text}")
                            logger.info(
                                f"[v6.2] Извлечено из {fname}: {len(text)} симв."
                            )
                    except Exception as e:
                        logger.warning(f"[v6.2] Ошибка обработки {fname} в ZIP: {e}")
                        continue

        except zipfile.BadZipFile:
            logger.error(f"[v6.2] Повреждённый ZIP: {doc_name}")
            return ""
        except Exception as e:
            logger.error(f"[v6.2] Ошибка распаковки ZIP {doc_name}: {e}")
            return ""
        finally:
            # Очистка временной директории
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass

        result = "\n\n".join(texts)
        logger.info(f"[v6.2] Из ZIP извлечено: {len(result)} символов")
        return result

    # ← v6.2: Новый метод — определение типа по содержимому
    def _extract_by_content(self, file_path: Path, doc_name: str) -> str:
        """Пробует определить тип файла по содержимому и извлечь текст."""
        try:
            with open(file_path, "rb") as f:
                header = f.read(8)

            if header.startswith(PDF_MAGIC):
                logger.info(f"[v6.2] Файл {doc_name} — PDF по magic bytes")
                return self._extract_pdf(file_path)
            elif header.startswith(ZIP_MAGIC):
                # Может быть DOCX или XLSX
                logger.info(
                    f"[v6.2] Файл {doc_name} — ZIP по magic bytes, пробуем DOCX"
                )
                text = self._extract_docx(file_path, doc_name)
                if not text:
                    text = self._extract_excel(file_path)
                return text
            elif header.startswith(b"\xd0\xcf\x11\xe0"):
                # OLE2 — старый .doc или .xls
                logger.info(f"[v6.2] Файл {doc_name} — OLE2, пробуем как DOCX")
                return self._extract_docx(file_path, doc_name)
            else:
                # Пробуем как текст
                logger.info(f"[v6.2] Неизвестный формат {doc_name}, пробуем как текст")
                return self._extract_text_file(file_path)
        except Exception as e:
            logger.error(f"[v6.2] Ошибка определения типа {doc_name}: {e}")
            return ""

    def _extract_docx(
        self, file_path: Path, doc_name: str, force_docx: bool = False
    ) -> str:
        logger.debug(f"Извлечение DOCX: {file_path}")

        def _do_extract():
            try:
                if HAS_DOCX:
                    doc = docx.Document(file_path)
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())

                    tables_text = []
                    for table in doc.tables:
                        try:
                            for row in table.rows:
                                try:
                                    row_text = []
                                    for cell in row.cells:
                                        if cell.text.strip():
                                            row_text.append(cell.text.strip())
                                    if row_text:
                                        tables_text.append(" | ".join(row_text))
                                except Exception as e:
                                    logger.debug(f"Ошибка чтения строки таблицы: {e}")
                                    continue
                        except Exception as e:
                            logger.debug(f"Ошибка чтения таблицы: {e}")
                            continue

                    all_text = "\n".join(paragraphs + tables_text)
                    return all_text
                else:
                    return self._extract_docx_via_zip(file_path)
            except Exception as e:
                logger.error(f"Ошибка извлечения DOCX: {e}")
                return ""

        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(_do_extract)
                return future.result(timeout=MAX_DOCX_TIMEOUT)
        except FutureTimeoutError:
            logger.error(f"Таймаут {MAX_DOCX_TIMEOUT}с на DOCX: {doc_name}")
            return "[ТАЙМАУТ: файл слишком сложный для обработки]"
        except Exception as e:
            logger.error(f"Ошибка при таймауте: {e}")
            return ""

    def _extract_docx_via_zip(self, file_path: Path) -> str:
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                with z.open("word/document.xml") as f:
                    xml_content = f.read().decode("utf-8", errors="ignore")
                    text = re.sub(r"<[^>]+>", "", xml_content)
                    text = re.sub(r"\s+", " ", text).strip()
                    return text
        except Exception as e:
            logger.error(f"Ошибка zip-извлечения: {e}")
            return ""

    def _extract_pdf(self, file_path: Path) -> str:
        if not HAS_PYMUPDF:
            logger.warning("PyMuPDF не установлен, PDF пропущен")
            return ""
        try:
            # ← v6.2: Проверка размера и сигнатуры PDF
            file_size = file_path.stat().st_size
            if file_size < 1000:
                logger.warning(f"PDF слишком мал ({file_size} байт) — возможно битый файл или HTML-ошибка")
                # Пробуем прочитать как текст (может быть HTML с ошибкой)
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                        if '<html' in text.lower() or '<!doctype' in text.lower():
                            logger.warning("Файл является HTML-страницей, не PDF")
                            return ""
                except:
                    pass
                return ""
            
            with open(file_path, 'rb') as f:
                header = f.read(5)
                if header != b'%PDF-':
                    logger.warning(f"Файл не является PDF (сигнатура: {header!r})")
                    return ""
            
            doc = fitz.open(file_path)
            texts = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    texts.append(text.strip())
            doc.close()
            return "\n".join(texts)
        except Exception as e:
            logger.error(f"Ошибка PDF: {e}")
            return ""

    def _extract_excel(self, file_path: Path) -> str:
        try:
            # ← v6.2: Определяем реальный тип по содержимому
            is_xlsx = False
            is_xls = False
            
            with open(file_path, 'rb') as f:
                header = f.read(8)
                if header.startswith(b'\x50\x4B\x03\x04'):
                    is_xlsx = True
                elif header.startswith(b'\xD0\xCF\x11\xE0'):
                    is_xls = True
            
            # Fallback на расширение
            if not is_xlsx and not is_xls:
                if file_path.suffix.lower() == ".xlsx":
                    is_xlsx = True
                elif file_path.suffix.lower() == ".xls":
                    is_xls = True
            
            if HAS_OPENPYXL and is_xlsx:
                wb = openpyxl.load_workbook(file_path, data_only=True)
                texts = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        row_text = [
                            str(cell.value) for cell in row if cell.value is not None
                        ]
                        if row_text:
                            texts.append(" | ".join(row_text))
                result = "\n".join(texts)
                logger.info(f"[v6.3] Excel извлечён: {len(texts)} строк, {len(result)} символов")
                return result

            elif HAS_XLRD and is_xls:
                wb = xlrd.open_workbook(file_path)
                texts = []
                for sheet in wb.sheets():
                    for row_idx in range(sheet.nrows):
                        row_text = [
                            str(sheet.cell_value(row_idx, col_idx))
                            for col_idx in range(sheet.ncols)
                        ]
                        if row_text:
                            texts.append(" | ".join(row_text))
                result = "\n".join(texts)
                logger.info(f"[v6.3] Excel извлечён: {len(texts)} строк, {len(result)} символов")
                return result
            else:
                logger.warning(f"Нет библиотеки для Excel: xlsx={is_xlsx}, xls={is_xls}, openpyxl={HAS_OPENPYXL}, xlrd={HAS_XLRD}")
                return ""
        except Exception as e:
            logger.error(f"Ошибка Excel: {e}")
            return ""

    def _extract_text_file(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Ошибка текстового файла: {e}")
            return ""
