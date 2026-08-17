"""
core/extractors/docx_extractor.py
Единый экстрактор текста из DOCX/DOC файлов.

ИСПРАВЛЕНО (v6.9.0):
  - Добавлен fallback для .doc (OLE) файлов:
    antiword → textract → olefile → LibreOffice
"""

import re
import zipfile
import subprocess
import tempfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from typing import Optional
from loguru import logger

from core.config.document_config import MAX_DOCX_TIMEOUT, QUANTITY_COLUMN_KEYWORDS
from core.extractors.base_extractor import BaseExtractor

try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx не установлен, используем zipfile fallback")


class DocxExtractor(BaseExtractor):
    """Единый экстрактор текста из DOCX/DOC файлов."""

    SUPPORTED_EXTENSIONS = ["docx", "doc"]

    def extract(self, file_path: Path, doc_name: str = "", **kwargs) -> str:
        """Извлекает текст из DOCX/DOC-файла."""
        path = Path(file_path)
        name = doc_name or path.name

        if not path.exists():
            logger.error(f"[DocxExtractor] Файл не найден: {path}")
            return ""

        # Проверяем валидность DOCX
        is_docx = self._is_valid_docx(path)

        # v6.9.0: Fallback для старых .doc (OLE)
        if not is_docx and path.suffix.lower() == ".doc":
            text = self._extract_doc_ole(path, name)
            if text:
                return text
            logger.error(f"[DocxExtractor] Не удалось прочитать DOC: {name}")
            return ""

        if not is_docx:
            logger.error(f"[DocxExtractor] Файл не является валидным DOCX: {name}")
            return ""

        # Сначала пробуем python-docx с таймаутом
        text = self._extract_with_docx(path, name)
        if text:
            return text

        # Fallback на zipfile
        logger.info(f"[DocxExtractor] Fallback на zipfile для {name}")
        return self._extract_via_zip(path, name)

    def _is_valid_docx(self, file_path: Path) -> bool:
        """Проверяет, является ли файл валидным DOCX (ZIP с XML)."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                if "word/document.xml" in z.namelist():
                    return True
                if "[Content_Types].xml" in z.namelist():
                    content = z.read("[Content_Types].xml").decode(
                        "utf-8", errors="ignore"
                    )
                    if "wordprocessingml" in content:
                        return True
            return False
        except zipfile.BadZipFile:
            return False
        except Exception as e:
            logger.warning(f"[DocxExtractor] Ошибка проверки DOCX: {e}")
            return False

    # ==================== v6.9.0: OLE/DOC fallback ====================

    def _extract_doc_ole(self, file_path: Path, doc_name: str) -> str:
        """Fallback для старых .doc файлов (OLE format)."""
        text = ""

        # Попытка 1: antiword
        try:
            result = subprocess.run(
                ["antiword", str(file_path)], capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0 and result.stdout:
                text = result.stdout
                logger.info(f"[DocxExtractor v6.9.0] antiword: {len(text)} симв.")
                return text
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Попытка 2: textract
        try:
            import textract

            raw = textract.process(str(file_path))
            text = raw.decode("utf-8", errors="ignore")
            logger.info(f"[DocxExtractor v6.9.0] textract: {len(text)} симв.")
            return text
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"[DocxExtractor] textract ошибка: {e}")

        # Попытка 3: olefile (pure Python)
        try:
            import olefile

            ole = olefile.OleFileIO(file_path)
            if ole.exists("WordDocument"):
                stream = ole.openstream("WordDocument")
                data = stream.read()
                text = self._extract_text_from_ole_word(data)
                logger.info(f"[DocxExtractor v6.9.0] olefile: {len(text)} симв.")
                return text
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"[DocxExtractor] olefile ошибка: {e}")

        # Попытка 4: LibreOffice headless
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    [
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmpdir,
                        str(file_path),
                    ],
                    capture_output=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    converted = Path(tmpdir) / file_path.with_suffix(".docx").name
                    if converted.exists():
                        return self.extract(converted, doc_name)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        return text

    def _extract_text_from_ole_word(self, data: bytes) -> str:
        """Простой извлекатель текста из WordDocument stream."""
        text_parts = []
        i = 0
        while i < len(data) - 1:
            if 32 <= data[i] <= 126 or 0xC0 <= data[i] <= 0xFF:
                try:
                    text_parts.append(
                        bytes([data[i]]).decode("cp1251", errors="ignore")
                    )
                except:
                    pass
            i += 1
        return "".join(text_parts)

    # ==================== DOCX extraction ====================

    def _extract_with_docx(self, file_path: Path, doc_name: str) -> str:
        """Извлекает текст через python-docx с таймаутом."""
        if not HAS_DOCX:
            return ""

        def _do_extract():
            try:
                document = docx.Document(file_path)
                paragraphs = []
                for para in document.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())

                tables_text = self._extract_tables(document)
                all_text = "\n".join(paragraphs + tables_text)
                return all_text
            except Exception as e:
                logger.error(f"[DocxExtractor] Ошибка python-docx: {e}")
                return ""

        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(_do_extract)
                return future.result(timeout=MAX_DOCX_TIMEOUT)
        except FutureTimeoutError:
            logger.error(
                f"[DocxExtractor] Таймаут {MAX_DOCX_TIMEOUT}с на {doc_name}, "
                f"переключаемся на zipfile"
            )
            return ""
        except Exception as e:
            logger.error(f"[DocxExtractor] Ошибка при таймауте: {e}")
            return ""

    def _extract_tables(self, document) -> list:
        """Извлекает таблицы с заголовками колонок."""
        tables_text = []

        for table_idx, table in enumerate(document.tables):
            try:
                rows_data = []
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        cell_text = self._extract_cell_text(cell)
                        cells.append(cell_text)
                    rows_data.append(cells)

                is_data_table = self._is_data_table(rows_data)

                if is_data_table:
                    formatted = self._format_data_table(rows_data, table_idx)
                    if formatted:
                        tables_text.append(formatted)
                else:
                    headers = self._extract_headers(table)
                    rows_text = self._extract_rows(table, headers)
                    tables_text.extend(rows_text)

            except Exception as e:
                logger.debug(f"[DocxExtractor] Ошибка таблицы: {e}")
                continue

        return tables_text

    def _extract_cell_text(self, cell) -> str:
        """Извлекает текст из ячейки, включая вложенные таблицы."""
        texts = []
        for para in cell.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        for nested_table in cell.tables:
            nested_text = self._format_table(nested_table, -1)
            if nested_text:
                texts.append(nested_text)
        return " ".join(texts)

    def _is_data_table(self, rows_data: list) -> bool:
        """Определяет, содержит ли таблица структурированные данные."""
        if len(rows_data) < 2:
            return False

        header = rows_data[0]
        data_keywords = [
            "№",
            "наименование",
            "количество",
            "кол-во",
            "ед.",
            "цена",
            "сумма",
            "должность",
            "рабочее место",
            "адрес",
            "кол-во",
        ]
        header_text = " ".join(h.lower() for h in header)
        has_header = any(kw in header_text for kw in data_keywords)

        numeric_count = 0
        total_cells = 0
        for row in rows_data[1:]:
            for cell in row:
                total_cells += 1
                if cell and re.search(r"\d+", cell):
                    numeric_count += 1

        has_numbers = numeric_count > 0 and (numeric_count / max(total_cells, 1)) > 0.2
        return has_header or has_numbers

    def _format_table(self, table, table_index: int) -> str:
        """Форматирует таблицу в читаемый вид (legacy для вложенных)."""
        rows_data = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = self._extract_cell_text(cell)
                cells.append(cell_text)
            rows_data.append(cells)

        if not rows_data:
            return ""

        is_data_table = self._is_data_table(rows_data)
        if is_data_table:
            return self._format_data_table(rows_data, table_index)
        else:
            return "\n".join(" | ".join(c for c in row if c) for row in rows_data)

    def _format_data_table(self, rows_data: list, table_index: int) -> str:
        """Форматирует таблицу данных в читаемый вид."""
        lines = [f"=== ТАБЛИЦА {table_index + 1} ==="]

        if rows_data:
            header = [h.strip() if h else "" for h in rows_data[0]]
            lines.append(" | ".join(header))
            lines.append("-" * (sum(len(h) for h in header) + 3 * len(header)))

        for row in rows_data[1:]:
            cells = [c.strip() if c else "" for c in row]
            if any(cells):
                lines.append(" | ".join(cells))

        lines.append("=== КОНЕЦ ТАБЛИЦЫ ===")
        return "\n".join(lines)

    def _extract_headers(self, table) -> list:
        """Извлекает заголовки таблицы (первая строка)."""
        headers = []
        try:
            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    headers.append(cell.text.strip().lower())
        except Exception as e:
            logger.debug(f"[DocxExtractor] Ошибка чтения заголовков: {e}")
        return headers

    def _extract_rows(self, table, headers: list) -> list:
        """Извлекает строки таблицы с подписями заголовков."""
        rows_text = []

        for row_idx, row in enumerate(table.rows):
            if row_idx == 0 and headers:
                continue

            try:
                row_text = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue

                    if col_idx < len(headers) and headers[col_idx]:
                        header = headers[col_idx]
                        if self._looks_like_quantity(header, cell_text):
                            row_text.append(f"{header}: {cell_text}")
                        else:
                            row_text.append(cell_text)
                    else:
                        row_text.append(cell_text)

                if row_text:
                    rows_text.append(" | ".join(row_text))
            except Exception as e:
                logger.debug(f"[DocxExtractor] Ошибка чтения строки: {e}")
                continue

        return rows_text

    def _looks_like_quantity(self, header: str, cell_text: str) -> bool:
        """Проверяет, похоже ли содержимое на количество по заголовку."""
        if not cell_text:
            return False
        if not re.match(r"^[\d\s.,]+$", cell_text.replace(" ", "")):
            return False
        header_lower = header.lower()
        return any(kw in header_lower for kw in QUANTITY_COLUMN_KEYWORDS)

    def _extract_via_zip(self, file_path: Path, doc_name: str) -> str:
        """Fallback: извлечение через zipfile (без python-docx)."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                if "word/document.xml" not in z.namelist():
                    logger.error(f"[DocxExtractor] Нет word/document.xml в архиве")
                    return ""

                xml_content = z.read("word/document.xml").decode(
                    "utf-8", errors="ignore"
                )

                text = re.sub(r"<[^>]+>", " ", xml_content)
                text = re.sub(r"\s+", " ", text)
                text = re.sub(r" (\d+\.) ", r"\n\1 ", text)

                result = text.strip()
                logger.info(
                    f"[DocxExtractor] Извлечено через zipfile: {len(result)} симв."
                )
                return result

        except Exception as e:
            logger.error(f"[DocxExtractor] Ошибка zip-извлечения: {e}")
            return ""
